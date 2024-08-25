import easyocr
import os
import fitz
from utils.log import get_logger
from utils.config import config
from algo.llm_inference import gpt_manager
from algo.llm_base_model import Prompt
import time
import datetime
import requests
import sys
import docx
import traceback
from dao.tool_dao import *
from io import StringIO
import csv
import codecs
from flask_admin._compat import csv_encode
from dao.task_dao import get_chats_by_job_id_with_date, query_candidate_by_id
from dao.manage_dao import get_job_name_by_id
from dao.contact_bank_dao import *
import json5
import re
from os.path import basename
from service.extension_service import refresh_contact
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from dao.contact_bank_dao import query_user_link_by_id_set, query_contact_by_id_set

from algo.llm_inference import gpt_manager
from algo.llm_base_model import Prompt

logger = get_logger(config['log']['log_file'])
reader = easyocr.Reader(['ch_sim', 'en'])  # this needs to run only once to load the model into memory

file_path_prefix = '/home/human/workspace/hunter_robot.v2.0/tmp/'

SCENARIO_GREETING = 'greeting'
SCENARIO_CHAT = 'chat'
SCENARIO_EMAIL = 'email'
SCENARIO_INMAIL = 'inmail'


def cv_str(obj, dent):
    cv = ""
    if type(obj) == dict:
        for k in obj:
            if obj[k]:
                for _ in range(dent):
                    cv += '\t'
                cv += (k + ":")
                cv += cv_str(obj[k], dent + 1)
    elif type(obj) == list:
        for e in obj:
            cv += cv_str(e, dent + 1)
    elif type(obj) == str or type(obj) == int or type(obj) == float:
        cv += (str(obj) + '\n')
    return cv


def deserialize_raw_profile(raw_profile):
    while type(raw_profile) == tuple:
        raw_profile = raw_profile[0]
    if raw_profile is None or type(raw_profile) != str:
        logger.error("[deserialize_raw_profile] raw profile not str")
        return None
    pattern = re.compile(r'•\s+')
    new_raw_profile = pattern.sub(' ', raw_profile)
    try:
        new_raw_profile = new_raw_profile.replace('\n', '\\n')
        if new_raw_profile.endswith('\\n'):
            new_raw_profile = new_raw_profile[:-2]
        return json.loads(new_raw_profile, strict=False)
    except BaseException as e:
        logger.error(f"deserialize_raw_profile error: {new_raw_profile}")
        logger.error(traceback.format_exc())
        return None


def get_age(profile):
    # 优先看本科开始年份 + 18
    # 如果有多个本科学历，选择开始时间最早的那个
    # 如果没有本科学历
    # 计算研究生学历，开始时间 + 21
    # 如果有多个研究生学历，选择开始时间最早的那个
    # 计算最早工作年限 + 21
    # 研究生学历计算出的年纪 和 工作年限计算出的年纪 谁大就选谁
    try:
        has_education = False
        has_master = False
        has_bachelor = False
        has_experience = False
        min_education_start_year = 1000000
        min_work_start_year = 1000000
        if 'educations' in profile or len(profile['educations']) > 0:
            for education in profile['educations']:
                if 'Bachelor' in education['degreeInfo'] or 'bachelor' in education['degreeInfo'] or 'Bachelor' in \
                        education['majorInfo'] or 'bachelor' in education['majorInfo']:
                    has_bachelor = True
                    min_education_start_year = min(get_min_time_info(education['timeInfo'], min_education_start_year),
                                                   min_education_start_year)
                elif 'Master' in education['degreeInfo'] or 'master' in education['degreeInfo'] or 'Master' in \
                        education['majorInfo'] or 'master' in education['majorInfo']:
                    has_master = True
                    min_education_start_year = min(get_min_time_info(education['timeInfo'], min_education_start_year),
                                                   min_education_start_year)
                else:
                    min_education_start_year = min(get_min_time_info(education['timeInfo'], min_education_start_year),
                                                   min_education_start_year)
                has_education = True
        # logger.info(f"1 get_age => min_education_start_year: {min_education_start_year}")
        if 'experiences' in profile and len(profile['experiences']) > 0:
            for experience in profile['experiences']:
                if 'companyName' in experience['companyName'] and (
                        'intern' in experience['companyName'] or 'Intern' in experience['companyName']):
                    continue
                intern = False
                if 'works' in experience and len(experience['works']) > 0:
                    for work in experience['works']:
                        if 'workPosition' in work and (
                                'intern' in work['workPosition'] or 'Intern' in work['workPosition']):
                            continue
                        if "workTimeInfo" in work:
                            min_work_start_year = min(get_min_time_info(work['workTimeInfo'], min_work_start_year), min_work_start_year)

                if intern:
                    continue
                has_experience = True
                min_work_start_year = min(get_min_time_info(experience['timeInfo'], min_work_start_year),
                                          min_work_start_year)
                # logger.info(
                #     f"get_age experiences timeInfo => {experience['timeInfo']} min_work_start_year: {min_work_start_year}")
        # logger.info(f"2 get_age => min_work_start_year: {min_work_start_year}")
        age_sure = None
        age_compare = None
        if has_education and has_bachelor and min_education_start_year < 1000000:
            # print('has_education has_bachelor')
            age_sure = 18 + datetime.datetime.now().year - min_education_start_year
        elif has_master and min_education_start_year < 1000000:
            # print('has_education has_master')
            age_compare = 21 + datetime.datetime.now().year - min_education_start_year
        elif has_education and min_education_start_year < 1000000:
            # print('has_education')
            age_sure = 18 + datetime.datetime.now().year - min_education_start_year
        if age_sure:
            # print("age_sure = ...")
            # logger.info(
            #     f"get_age age_sure: {age_sure} has_education: {has_education} has_bachelor: {has_bachelor} min_education_start_year: {min_education_start_year}")
            return None if age_sure > 100 or age_sure <= 0 else age_sure

        if not has_experience:
            # print('age compare !has_experience')
            return None if age_compare > 100 else age_compare
        elif age_compare is None:
            # print('has_experience age_compar is none')
            return None if min_work_start_year == 1000000 else 21 + datetime.datetime.now().year - min_work_start_year
        else:
            # print('has_experience age_compare max')
            age = max(age_compare, 21 + datetime.datetime.now().year - min_work_start_year)
            # logger.info(f"get_age age: {age_sure}")
            return None if age > 100 or age <= 0 else age
    except BaseException as e:
        logger.error(f"get_age error => {e}")
        return None


def cal_work_time(experiences):
    min_work_start_year = 2900

    for experience in experiences:
        if 'companyName' in experience['companyName'] and (
                'intern' in experience['companyName'] or 'Intern' in experience['companyName']):
            continue
        intern = False
        if 'works' in experience and len(experience['works']) > 0:
            for work in experience['works']:
                if 'workPosition' in work and (
                        'intern' in work['workPosition'] or 'Intern' in work['workPosition']):
                    continue
                if "workTimeInfo" in work:
                    min_work_start_year = min(get_min_time_info(work['workTimeInfo'], min_work_start_year),
                                              min_work_start_year)
        if intern:
            continue
        # logger.info(f"cal_work_time {experience['timeInfo']}")
        min_work_start_year = min(get_min_time_info(experience['timeInfo'], 100000),
                                  min_work_start_year)
    # logger.info(f"cal_work_time min_work_start_year: {min_work_start_year}")
    return None if min_work_start_year == 2900 else min_work_start_year


def cal_company(companyName):
    company = companyName.split(' · ')[0]
    return company


def parse_profile(profile, type='need_deserialize', field_2_str=False):
    if type == 'need_deserialize':
        profile = deserialize_raw_profile(profile)
    if profile is None:
        return None
    res = {'candidateId': None,
           'department': None,
           'company': None,
           'title': None,
           'lastTitle': None,
           'last5Jump': None,
           'name': None,
           'location': None,
           'contactInfo': None,
           'cv': None,
           'age': None,
           'isChinese': None,
           'languages': None,
           'workTime': None}
    if 'id' in profile:
        res['candidateId'] = profile['id']
    if 'profile' in profile:
        profile = profile['profile']
    experience = profile['experiences'][0] if 'experiences' in profile and len(profile['experiences']) > 0 else None
    if experience and 'companyName' in experience:
        res['department'] = experience['companyName']
        res['company'] = cal_company(experience['companyName'])

    if 'role' in profile:
        res['title'] = profile['role']
    elif 'works' in experience and len(experience['works']) > 0:
        work = experience['works'][0]
        if 'workPosition' in work:
            res['role'] = work['workPosition']
    res['lastTitle'] = res['title']
    name = None
    if 'name' in profile:
        name = profile['name']
        res['name'] = profile['name']

    if 'location' in profile:
        res['location'] = profile['location']

    if 'contactInfo' in profile:
        res['contactInfo'] = cv_str(profile['contactInfo'], 0) if field_2_str else profile['contactInfo']
    if profile:
        res['cv'] = cv_str(profile, 0)
    if name:
        res['isChinese'] = False
        chs_names = ['Zhao', 'Qian', 'Sun', 'Li', 'Zhou', 'Wu', 'Zheng', 'Wang', 'Feng', 'Chen', 'Zhu', 'Wei', 'Shen',
                     'Han', 'Yang', 'Qin', 'You', 'Xu', 'He', 'Lv', 'Shi', 'Zhang', 'Kong', 'Cao', 'Yan', 'Hua', 'Jin',
                     'Tao', 'Jiang', 'Xie', 'Zou', 'Yu', 'Bo', 'Shui', 'Dou', 'Yun', 'SU', 'Pan', 'Ge', 'Fan', 'Peng',
                     'Lang', 'Lu', 'Chang', 'Ma', 'Miao', 'Feng', 'Hua', 'Fang', 'Yu', 'Ren', 'Yuan', 'Liu', 'Bao',
                     'Shi', 'Tang', 'Fei', 'Lian', 'Qin', 'Xue', 'Lei', 'He', 'Ni', 'Teng', 'Yin', 'Luo', 'Bi', 'Hao',
                     'Wu', 'An', 'Chang', 'Le', 'Yu', 'Fu', 'Pi', 'Qi', 'Kang', 'Bu', 'Gu', 'Meng', 'Ping', 'Huang',
                     'He', 'Mu', 'Xiao', 'Yin', 'Yao', 'Shao', 'Qi', 'Mao', 'Di', 'Mi', 'Bei', 'Ming', 'Zang', 'Ji',
                     'FU', 'Cheng', 'Dai', 'Song', 'Ji', 'Shu', 'Qu', 'Dong', 'Liang', 'Du', 'Lan', 'Min', 'Jia', 'Lou',
                     'Tong', 'Guo', 'Lin', 'Diao', 'Zhong', 'Qiu', 'Luo', 'Gao', 'Xia', 'Cai', 'Tian', 'Hu', 'Ling',
                     'Huo', 'Ling', 'Wan', 'Zhi', 'ke', 'Guan', 'Mo', 'Miao', 'Xie', 'Zong', 'Ding', 'Deng', 'Shan',
                     'Hang', 'Bao', 'Zuo', 'Cui', 'Niu', 'Weng', 'Xun', 'Yang', 'Hui', 'Gong', 'Cheng', 'Hua', 'Pei',
                     'Rong', 'Jiao', 'Mu', 'Gu', 'Che', 'Hou', 'Mi', 'Quan', 'Ban', 'Gong', 'Ning', 'Chou', 'Luan',
                     'Zu', 'Fu', 'Liu', 'Long', 'Ye', 'Si', 'Bai', 'Huai', 'Cong', 'Lai', 'Zhuo', 'Qiao', 'Shuang',
                     'Dang', 'Cui', 'Tan', 'Ran', 'Bian', 'Chai', 'Liao', 'Gong', 'Jian', 'Sha', 'You', 'Hai', 'Wen',
                     'Zhai', 'Kou', 'Rao', 'Pu', 'Ou', 'She', 'Nian', 'Ai', 'Ha', 'An', 'Zhan', 'Ruan', 'Bing', 'Tu',
                     'Zhuang', 'Geng', 'Guang', 'Chao', 'AH', 'AU', 'BIK', 'BING', 'BIT', 'BONG', 'BUN', 'CHAI', 'CHAK',
                     'CHAM', 'CHAN', 'CHANG', 'CHAT', 'CHAU', 'CHEN', 'CHENG', 'CHEONG', 'CHEUK', 'CHEUNG', 'CHI',
                     'CHIANG', 'CHICKC', 'HIGN', 'CHIK', 'CHIN', 'CHING', 'CHIT', 'CHIU', 'CHO', 'CHOI', 'CHOK',
                     'CHONG', 'CHOR', 'CHOW', 'CHOY', 'CHU', 'CHUEN', 'CHUI', 'CHUM', 'CHUN', 'CHUNG', 'DIK', 'DIU',
                     'FAT', 'FA', 'FAI', 'FAN', 'FANG', 'FEI', 'FO', 'FOG', 'FOK', 'FONG', 'FOO', 'FOOK', 'FOON',
                     'FORK', 'FU', 'FUI', 'FUK', 'FUNG', 'HING', 'HA', 'HAN', 'HANG', 'HAU', 'HEI', 'HEUNG', 'HIM',
                     'HIN', 'HIP', 'HIU', 'HO', 'HOHO', 'HOI', 'HOK', 'HON', 'HONG', 'HOU', 'HSU', 'HSUI', 'HUANG',
                     'HUEN', 'HUI', 'HUNG', 'HWANG', 'JIM', 'KA', 'KAI', 'KAK', 'KAM', 'KAN', 'KANG', 'KAR', 'KAU',
                     'KEI', 'KEUNG', 'KHOO', 'KIM', 'KIN', 'KING', 'KIT', 'KIU', 'KO', 'KOK', 'KON', 'KONG', 'KOON',
                     'KOT', 'KU', 'KUA', 'KUEN', 'KUI', 'KUK', 'KUN', 'KUNG', 'KUO', 'KWAI', 'KWAN', 'KWING', 'KWOK',
                     'KWONG', 'LAI', 'LAM', 'LAN', 'LAP', 'LARM', 'LAU', 'LAW', 'LEE', 'LEI', 'LEONG', 'LEUNG', 'LI',
                     'LIANG', 'LIAO', 'LIEW', 'LIK', 'LIM', 'LIN', 'LING', 'LIP', 'LIT', 'LIU', 'LO', 'LOI', 'LOK',
                     'LONG', 'LOO', 'LOOK', 'LOONG', 'LOW', 'LUEN', 'LUET', 'LUI', 'LUK', 'LUMLUN', 'LUN', 'LUNG', 'MA',
                     'MAK', 'MAN', 'MANG', 'MAO', 'MAR', 'MEI', 'MIN', 'MING', 'MIU', 'MO', 'MOK', 'MOOK', 'MOON',
                     'MUI', 'MUK', 'MUNG', 'NAM', 'NANG', 'NAR', 'NEI', 'NEUNG', 'NG', 'NGA', 'NGAI', 'NGAN', 'NGAU',
                     'NGO', 'NGON', 'NIE', 'NIN', 'NING', 'NUI', 'O', 'OI', 'ON', 'PAK', 'PANG', 'PAT', 'PAU', 'PEI',
                     'PIK', 'PIN', 'PING', 'PIU', 'PO', 'POK', 'PONG', 'POO', 'POON', 'PUI', 'PUN', 'SAI', 'SAM', 'SAN',
                     'SANG', 'SAU', 'SE', 'SECK', 'SEE', 'SEI', 'SEK', 'SHAN', 'SHE', 'SHEK', 'SHEUNG', 'SHI', 'SHIH',
                     'SHING', 'SHIU', 'SHP', 'SHU', 'SHUE', 'SHUEN', 'SHUK', 'SHUM', 'SHUN', 'SI', 'SIK', 'SIM', 'SIN',
                     'SING', 'SIT', 'SIU', 'SO', 'SUEN', 'SUET', 'SUI', 'SUM', 'SUN', 'SUNG', 'SZE', 'TAI', 'TAK',
                     'TAM', 'TAN', 'TANG', 'TAO', 'TAT', 'TAU', 'TIM', 'TIN', 'TING', 'TIP', 'TIT', 'TO', 'TONG',
                     'TSAM', 'TSANG', 'TSE', 'TSIM', 'TSO', 'TSOI', 'TSUI', 'TUEN', 'TUNG', 'TYE', 'UNG', 'VONG', 'WAH',
                     'WAI', 'WAN', 'WANG', 'WAT', 'WING', 'WO', 'WON', 'WONG', 'WOO', 'WOOD', 'WOON', 'WU', 'WUI',
                     'WUN', 'WUT', 'YAM', 'YAN', 'YANG', 'YAO', 'YAT', 'YAU', 'YEE', 'YEI', 'YEN', 'YEUK', 'YEUNG',
                     'YI', 'YICK', 'YIK', 'YIM', 'YIN', 'YING', 'YIP', 'YIU', 'YOUNG', 'YU', 'YUE', 'YUEN', 'YUET',
                     'YUI', 'YUK', 'YUNG', 'ZHANG']
        for split_name in name.split(' '):
            if ('\u4E00' <= split_name <= '\u9FFF') or ('\u3400' <= split_name <= '\u4DBF') or split_name in chs_names:
                res['isChinese'] = True
    # age
    res['age'] = get_age(profile)
    # logger.info(f"get_age => {res['age']}")
    if 'languages' in profile and len(profile['languages']) > 0:
        res['languages'] = cv_str(profile['languages'], 0) if field_2_str else profile['languages']

    if 'experiences' in profile and len(profile['experiences']) > 0:
        experiences = profile['experiences']
        last_5_jump = 0
        start_year = datetime.datetime.now().year - 5
        for experience in experiences:
            # if 'timeInfo' in experience and experience['timeInfo'] != None and type(experience['timeInfo']) == str:
            if 'timeInfo' in experience and experience['timeInfo'] != None:
                if (get_max_time_info(experience['timeInfo'], 1000)) > start_year:
                    last_5_jump += 1
        res['last5Jump'] = last_5_jump
        res["experiences"] = experiences
        res["workTime"] = cal_work_time(experiences)
    # logger.info(f"age: {res['age']}")
    return res


def get_candidate_id(profile, platform):
    if platform == 'maimai':
        return profile['id']
    if platform == 'Linkedin':
        return profile['id']
    if platform == 'Boss':
        return profile['geekCard']['geekId']
    if platform == 'liepin':
        return profile['usercIdEncode']


def maimai_online_resume_upload_processor(manage_account_id, profile, platform, tag):
    if tag and len(tag) > 0:
        create_profile_tag(manage_account_id, platform, tag)
    count = 0
    for p in profile:
        candidate_id = get_candidate_id(p, platform)
        if candidate_id == None or candidate_id == '':
            continue
        if len(get_resume_by_candidate_id_and_platform(candidate_id, platform, manage_account_id)) == 0:
            exp = []
            for e in p.get('exp', []):
                des = e["description"] or ''
                des = des.replace('"', "").replace("'", "").replace("\n", ";").replace('\"', "").replace("\'", "")
                exp.append({
                    "company": e["company"],
                    "v": e["v"],
                    "position": e["position"],
                    "worktime": e["worktime"],
                    "description": des
                })
            p['exp'] = exp
            parsed = parse_profile(p, 'no_need')
            upload_online_profile(manage_account_id, platform, json.dumps(p, ensure_ascii=False), candidate_id,
                                  parsed['name'] if parsed['name'] else '',
                                  parsed['company'] if parsed['company'] else '')
            count = count + 1
        if tag and len(tag) > 0:
            associate_profile_tags(manage_account_id, candidate_id, platform, tag)
    return count


def linkedin_filter(manage_account_id, raw_profile, conditions, platform):
    linkedin_online_resume_upload_processor(manage_account_id, [raw_profile], platform, '', -20000, 20000, '')
    profile = parse_profile(raw_profile, 'no_need')

    if "age" in conditions.keys():
        if profile["age"] is None:
            return True

        if profile["age"] < conditions["age"]["min_age"] or profile["age"] > conditions["age"]["max_age"]:
            return True

    if "is_chinese" in conditions.keys() and conditions["is_chinese"]:
        if not profile["isChinese"]:
            return True

    return False


def linkedin_online_resume_upload_processor(manage_account_id, profile, platform, list_name, min_age, max_age, tag):
    if tag and len(tag) > 0:
        create_profile_tag(manage_account_id, platform, tag)
    count = 0
    for p in profile:
        candidate_id = get_candidate_id(p, platform)
        if list_name != '':
            add_list_relation(manage_account_id, list_name, candidate_id)
        if candidate_id == None or candidate_id == '':
            continue
        firt_work_year = 10000
        current_year = int(datetime.datetime.now().year)
        if len(get_resume_by_candidate_id_and_platform(candidate_id, platform,
                                                       manage_account_id)) == 0 and 'profile' in p:
            for l in p.get('profile', {}).get('languages', []):
                language = l.get('language', '') or ''
                l['language'] = language.replace('"', "").replace("'", "").replace("\n", ";").replace('\"', "").replace(
                    "\'", "")
                des = l.get('des', '') or ''
                l['des'] = des.replace('"', "").replace("'", "").replace("\n", ";").replace('\"', "").replace("\'", "")
            for e in p.get('profile', {}).get('experiences', []):
                companyName = e.get('companyName', '') or ''
                e['companyName'] = companyName.replace('"', "").replace("'", "").replace("\n", ";").replace('\"',
                                                                                                            "").replace(
                    "\'", "")
                for w in e.get('works', []):
                    workTimeInfo = w.get('workTimeInfo', '') or ''
                    w['workTimeInfo'] = workTimeInfo.replace('"', "").replace("'", "").replace("\n", ";").replace('\"',
                                                                                                                  "").replace(
                        "\'", "")
                    # 截年龄
                    years = re.findall(r'\b\d{4}\b', w['workTimeInfo'])
                    min_year = 10000 if len(years) == 0 else int(min(years))
                    if min_year < firt_work_year:
                        firt_work_year = min_year
                    workLocationInfo = w.get('workLocationInfo', '') or ''
                    w['workLocationInfo'] = workLocationInfo.replace('"', "").replace("'", "").replace("\n",
                                                                                                       ";").replace(
                        '\"', "").replace("\'", "")
                    workPosition = w.get('workPosition', '') or ''
                    w['workPosition'] = workPosition.replace('"', "").replace("'", "").replace("\n", ";").replace('\"',
                                                                                                                  "").replace(
                        "\'", "")
                    workDescription = w.get('workDescription', '') or ''
                    w['workDescription'] = workDescription.replace('"', "").replace("'", "").replace("\n", ";").replace(
                        '\"', "").replace("\'", "")
            for edu in p.get('profile', {}).get('educations', []):
                summary = edu.get('summary', '') or ''
                edu['summary'] = summary.replace('"', "").replace("'", "").replace("\n", ";").replace('\"', "").replace(
                    "\'", "")
                degreeInfo = edu.get('degreeInfo', '') or ''
                edu['degreeInfo'] = degreeInfo.replace('"', "").replace("'", "").replace("\n", ";").replace('\"',
                                                                                                            "").replace(
                    "\'", "")
                majorInfo = edu.get('majorInfo', '') or ''
                edu['majorInfo'] = majorInfo.replace('"', "").replace("'", "").replace("\n", ";").replace('\"',
                                                                                                          "").replace(
                    "\'", "")
                timeInfo = edu.get('timeInfo', '') or ''
                edu['timeInfo'] = timeInfo.replace('"', "").replace("'", "").replace("\n", ";").replace('\"',
                                                                                                        "").replace(
                    "\'", "")
                # 截年龄
                years = re.findall(r'\b\d{4}\b', edu['timeInfo'])
                max_year = 10000 if len(years) == 0 else int(max(years))
                if max_year < firt_work_year:
                    firt_work_year = max_year
                schoolName = edu.get('schoolName', '') or ''
                edu['schoolName'] = schoolName.replace('"', "").replace("'", "").replace("\n", ";").replace('\"',
                                                                                                            "").replace(
                    "\'", "")
            age = current_year - firt_work_year + 23
            if min_age > age or max_age < age:
                logger.info(f'profile_age_filter：{manage_account_id}, {candidate_id}, {age}')
            else:
                count = count + 1

            summary = p.get('profile', {}).get('summary', '') or ''
            role = p.get('profile', {}).get('role', '') or ''
            location = p.get('profile', {}).get('location', '') or ''
            name = p.get('profile', {}).get('name', '') or ''

            url = p.get('profile', {}).get('contactInfo', {}).get("url", "") or ''
            p['profile']['contactInfo']["url"] = url.replace('"', "").replace("'", "").replace("\n", ";").replace('\"',
                                                                                                                  "").replace(
                "\'", "")
            phone = p.get('profile', {}).get('contactInfo', {}).get("Phone", "") or ''
            p['profile']['contactInfo']["Phone"] = phone.replace('"', "").replace("'", "").replace("\n", ";").replace(
                '\"', "").replace("\'", "")
            email = p.get('profile', {}).get('contactInfo', {}).get("Email", "") or ''
            p['profile']['contactInfo']["Email"] = email.replace('"', "").replace("'", "").replace("\n", ";").replace(
                '\"', "").replace("\'", "")

            p['profile']['summary'] = summary.replace('"', "").replace("'", "").replace("\n", ";").replace('\"',
                                                                                                           "").replace(
                "\'", "")
            p['profile']['role'] = role.replace('"', "").replace("'", "").replace("\n", ";").replace('\"', "").replace(
                "\'", "")
            p['profile']['location'] = location.replace('"', "").replace("'", "").replace("\n", ";").replace('\"',
                                                                                                             "").replace(
                "\'", "")
            p['profile']['name'] = name.replace('"', "").replace("'", "").replace("\n", ";").replace('\"', "").replace(
                "\'", "")
            parsed = parse_profile(p, 'no_need')
            upload_online_profile(manage_account_id, platform, json.dumps(p, ensure_ascii=False), candidate_id,
                                  parsed['name'] if parsed['name'] else '',
                                  parsed['company'] if parsed['company'] else '')

        if tag and len(tag) > 0:
            associate_profile_tags(manage_account_id, candidate_id, platform, tag)

        try:
            upload_profile_status(manage_account_id, candidate_id, platform, p["profile"])
            refresh_contact(manage_account_id, candidate_id, p)
        except BaseException as e:
            logger.error("简历信息存储出错:", e)
            logger.error(str(traceback.format_exc()))

    return count


def filter_already_linkedin_ids(manage_account_id, linkedin_ids):
    rest_linkedin_ids = []
    for linkedin_id in linkedin_ids:
        status = query_profile_status_dao(manage_account_id, linkedin_id, 'Linkedin')
        if status == "connected":
            continue
        rest_linkedin_ids.append(linkedin_id)
    return rest_linkedin_ids


def generate_candidate_csv_by_job_liepin(job_id, start_date, end_date):
    chat_list = get_chats_by_job_id_with_date(job_id, start_date, end_date)
    job_name = get_job_name_by_id(job_id)
    io = StringIO()
    w = csv.writer(io)

    l = ['岗位名称', '候选人ID', '创建时间', '候选人姓名', '来源', '微信', '电话', '简历', '对话详情', '性别',
         '生日年份', '工作年限', '岗位', '学历', '地点', '薪资', '学校经历', '工作经历']
    l_encode = [csv_encode(_l) for _l in l]
    l_encode[0] = codecs.BOM_UTF8.decode("utf8") + codecs.BOM_UTF8.decode() + l_encode[0]
    w.writerow(l_encode)
    yield io.getvalue()
    io.seek(0)
    io.truncate(0)
    for c in chat_list:
        try:
            candidate_info = query_candidate_by_id(c[2])
            job_name = job_name
            candidate_id = c[2]
            create_time = c[9].strftime("%Y-%m-%d %H:%M:%S")
            candidate_name = c[3]
            if c[4] == 'user_ask':
                source = '候选人主动'
            elif c[4] == 'search':
                source = '机器人打招呼'
            else:
                source = '未知'
            if c[6] is None:
                wechat = ''
                phone = ''
                resume = ''
            else:
                try:
                    contact = json.loads(c[6])
                    wechat = contact['wechat'] or ''
                    phone = contact['phone'] or ''
                    resume = contact['cv'] or ''
                except Exception as e:
                    wechat = ''
                    phone = ''
                    resume = ''
                    logger.info(
                        f'exception_filter:{candidate_id}, {candidate_name}, {c[6]}, {contact}, {e}, {e.args}, {traceback.format_exc()}')
            try:
                conversation = json.loads(c[7])
                con_str = ''
                for c in conversation:
                    con_str = con_str + c['speaker'] + ':' + c['msg'] + '\n'
            except Exception as e:
                con_str = ''

            if len(candidate_info) == 0:
                logger.info(f"chat_candidate_not_match, {candidate_id}")
                gender = ''
                born_year = ''
                work_year = ''
                position = ''
                degree = ''
                location = ''
                salary = ''
                edu = ''
                work = ''
            else:
                c_j = candidate_info[0][7].replace('\n', '.')
                c_j = c_j.replace("\'", '\"')
                candidate_json = json.loads(c_j, strict=False)
                gender = candidate_json.get('basicInfoForm', {}).get('sex', '')
                born_year = candidate_json.get('simpleResumeForm', {}).get('resBirthYear', '')
                work_year = candidate_json.get('basicInfoForm', {}).get('workYearsDescr', '')
                position = candidate_json.get('basicInfoForm', {}).get('resTitle', '')
                degree = candidate_json.get('basicInfoForm', {}).get('eduLevelName', '')
                location = candidate_json.get('basicInfoForm', {}).get('dqName', '')
                salary = candidate_json.get('basicInfoForm', {}).get('salary', '')
                edu = ''
                for e in candidate_json.get('eduExpFormList', []):
                    edu = edu + str(e.get('startYear', '')) + '-' + str(e.get('endYear', '')) + ', ' + str(
                        e.get('redDegreeName', '')) + ', ' + str(e.get('redSchool', '')) + ', ' + str(
                        e.get('redSpecial', '')) + '\n\n'
                work = ''
                for wo in candidate_json.get('workExps', []):
                    work = work + str(wo.get('startYear', '')) + '-' + str(wo.get('endYear', '')) + ', ' + str(
                        wo.get('rwCompname', '')) + ', ' + str(wo.get('rwDqName', '')) + '\n' + str(
                        wo.get('rwDuty', '')) + '\n\n'

            l = [job_id, candidate_id, create_time, candidate_name, source, wechat, phone, resume, con_str, gender,
                 born_year, work_year, position, degree, location, salary, edu, work]
            l_encode = [csv_encode(_l) for _l in l]
            w.writerow(l_encode)
            yield io.getvalue()
            io.seek(0)
            io.truncate(0)
        except Exception as e:
            logger.info(f'test_download_candidate_liepin_error4, {c_j}')
            logger.info(
                f'test_download_candidate_liepin_error4,{candidate_id}, {e}, {e.args}, {traceback.format_exc()}')


def generate_candidate_csv_by_job_Boss(job_id, start_date, end_date):
    chat_list = get_chats_by_job_id_with_date(job_id, start_date, end_date)
    job_name = get_job_name_by_id(job_id)
    io = StringIO()
    w = csv.writer(io)

    l = ['岗位名称', '候选人ID', '创建时间', '候选人姓名', '来源', '微信', '电话', '简历', '对话详情', '薪资范围',
         '年龄', '最高学历', '性别', '状态', '学校', '教育经历', '工作经历']
    l_encode = [csv_encode(_l) for _l in l]
    l_encode[0] = codecs.BOM_UTF8.decode("utf8") + codecs.BOM_UTF8.decode() + l_encode[0]
    w.writerow(l_encode)
    yield io.getvalue()
    io.seek(0)
    io.truncate(0)
    for c in chat_list:
        try:
            candidate_info = query_candidate_by_id(c[2])
            job_name = job_name
            candidate_id = c[2]
            create_time = c[9].strftime("%Y-%m-%d %H:%M:%S")
            candidate_name = c[3]
            if c[4] == 'user_ask':
                source = '候选人主动'
            elif c[4] == 'search':
                source = '机器人打招呼'
            else:
                source = '未知'
            if c[6] is None:
                wechat = ''
                phone = ''
                resume = ''
            else:
                try:
                    contact = json.loads(c[6])
                    wechat = contact['wechat'] or ''
                    phone = contact['phone'] or ''
                    resume = contact['cv'] or ''
                except Exception as e:
                    wechat = ''
                    phone = ''
                    resume = ''
                    logger.info(
                        f'exception_filter:{candidate_id}, {candidate_name}, {c[6]}, {contact}, {e}, {e.args}, {traceback.format_exc()}')
            try:
                conversation = json.loads(c[7])
                con_str = ''
                for c in conversation:
                    con_str = con_str + c['speaker'] + ':' + c['msg'] + '\n'
            except Exception as e:
                con_str = ''

            if len(candidate_info) == 0:
                logger.info(f"chat_candidate_not_match, {candidate_id}")
                salary = ''
                age = ''
                degree = ''
                gender = ''
                status = ''
                school = ''
                edu = ''
                work = ''
            else:
                c_j = candidate_info[0][7].replace('\n', '.')
                c_j = c_j.replace("\'", '\"')
                candidate_json = json.loads(c_j, strict=False)
                salary = candidate_json.get('geekCard', {}).get('salary', '')
                age = candidate_json.get('geekCard', {}).get('ageDesc', '')
                degree = candidate_json.get('geekCard', {}).get('geekDegree', '')
                gender = candidate_json.get('geekCard', {}).get('geekGender', '')
                status = candidate_json.get('geekCard', {}).get('applyStatusDesc', '')
                school = candidate_json.get('geekCard', {}).get('geekEdu', {}).get('school', '')
                edu = ''
                for s in candidate_json.get('geekCard', {}).get('geekEdus', []):
                    edu = edu + s.get('school', '') or '' + ',' + s.get('major', '') or '' + ',' + s.get('degreeName',
                                                                                                         '') or '' + '\n' + s.get(
                        'startDate', '') or '' + '-' + s.get('endDate', '') or '' + '\n\n'
                work = ''
                for e in candidate_json.get('geekCard', {}).get('geekWorks', []):
                    work = work + e.get('company', '') or '' + ',' + e.get('workTime', '') or '' + ',' + e.get(
                        'positionName', '') or '' + '\n'
                    work = work + e.get('startDate', '') or '' + '-' + e.get('endDate', '') or '' + '\n'
                    work = work + e.get('responsibility', '') or '' + '\n\n'

            l = [job_id, candidate_id, create_time, candidate_name, source, wechat, phone, resume, con_str, salary, age,
                 degree, gender, status, school, edu, work]
            l_encode = [csv_encode(_l) for _l in l]
            w.writerow(l_encode)
            yield io.getvalue()
            io.seek(0)
            io.truncate(0)
        except Exception as e:
            logger.info(f'test_download_candidate_boss_error4, {c_j}')
            logger.info(f'test_download_candidate_boss_error4,{candidate_id}, {e}, {e.args}, {traceback.format_exc()}')


def generate_candidate_csv_by_job_Linkedin(job_id, start_date, end_date):
    chat_list = get_chats_by_job_id_with_date(job_id, start_date, end_date)
    job_name = get_job_name_by_id(job_id)
    io = StringIO()
    w = csv.writer(io)

    l = ['岗位名称', '候选人ID', '创建时间', '候选人姓名', '来源', '微信', '电话', '邮箱', '简历', '对话详情', '地区',
         '岗位', '学校经历', '公司经历', '语言能力']
    l_encode = [csv_encode(_l) for _l in l]
    l_encode[0] = codecs.BOM_UTF8.decode("utf8") + codecs.BOM_UTF8.decode() + l_encode[0]
    w.writerow(l_encode)
    yield io.getvalue()
    io.seek(0)
    io.truncate(0)
    for c in chat_list:
        try:
            candidate_info = query_candidate_by_id(c[2])
            job_name = job_name
            candidate_id = c[2]
            create_time = c[9].strftime("%Y-%m-%d %H:%M:%S")
            candidate_name = c[3]
            if c[4] == 'user_ask':
                source = '候选人主动'
            elif c[4] == 'search':
                source = '机器人打招呼'
            else:
                source = '未知'
            if c[6] is None:
                wechat = ''
                phone = ''
                resume = ''
                email = ''
            else:
                try:
                    contact = json.loads(c[6])
                    wechat = contact['wechat'] or ''
                    phone = contact['phone'] or ''
                    resume = contact['cv'] or ''
                    email = contact['email'] or ''
                except Exception as e:
                    wechat = ''
                    phone = ''
                    resume = ''
                    email = ''
                    logger.info(
                        f'exception_filter:{candidate_id}, {candidate_name}, {c[6]}, {contact}, {e}, {e.args}, {traceback.format_exc()}')
            try:
                conversation = json.loads(c[7])
                con_str = ''
                for c in conversation:
                    con_str = con_str + c['speaker'] + ':' + c['msg'] + '\n'
            except Exception as e:
                con_str = ''

            if len(candidate_info) == 0:
                logger.info(f"chat_candidate_not_match, {candidate_id}")
                region = ''
                position = ''
                edu = ''
                work = ''
                language = ''
            else:
                c_j = candidate_info[0][7].replace('\n', '.')
                c_j = c_j.replace("\'", '\"')
                candidate_json = json.loads(c_j, strict=False)
                region = candidate_json.get('profile', {}).get('location', '')
                position = candidate_json.get('profile', {}).get('role', '')
                edu = ''
                for s in candidate_json.get('profile', {}).get('educations', []):
                    edu = edu + s.get('schoolName', '') + ',' + s.get('majorInfo', '') + ',' + s.get('degreeInfo',
                                                                                                     '') + '\n' + s.get(
                        'timeInfo', '') + '\n' + s.get('summary', '') + '\n\n'
                work = ''
                for e in candidate_json.get('profile', {}).get('experiences', []):
                    work = work + e.get('companyName', '') + ',' + e.get('timeInfo', '') + '\n'
                    for wo in e.get('works', []):
                        work = work + "positon:" + wo.get('workPosition', '') + '\n'
                        work = work + "time:" + wo.get('workTimeInfo', '') + '\n'
                        work = work + "location:" + wo.get('workLocationInfo', '') + '\n'
                        work = work + "description:" + wo.get('workDescription', '') + '\n\n'
                language = ''
                for lan in candidate_json.get('profile', {}).get('languages', []):
                    language = language + lan.get('language', '') + '\n' + lan.get('des', '') + '\n\n'

            l = [job_id, candidate_id, create_time, candidate_name, source, wechat, phone, email, resume, con_str,
                 region, position, edu, work, language]
            l_encode = [csv_encode(_l) for _l in l]
            w.writerow(l_encode)
            yield io.getvalue()
            io.seek(0)
            io.truncate(0)
        except Exception as e:
            logger.info(f'test_download_candidate_linkedin_error4, {c_j}')
            logger.info(
                f'test_download_candidate_linkedin_error4,{candidate_id}, {e}, {e.args}, {traceback.format_exc()}')


def generate_candidate_csv_by_job_maimai(job_id, start_date, end_date):
    chat_list = get_chats_by_job_id_with_date(job_id, start_date, end_date)
    job_name = get_job_name_by_id(job_id)
    io = StringIO()
    w = csv.writer(io)

    l = ['岗位名称', '候选人ID', '创建时间', '候选人姓名', '来源', '微信', '电话', '简历', '对话详情', '地区', '性别',
         '年龄', '岗位', '最高学历', '专业', '历史公司', '毕业院校', '教育经历', '工作经历', '预期地点', '预期薪水',
         '简历标签']
    l_encode = [csv_encode(_l) for _l in l]
    l_encode[0] = codecs.BOM_UTF8.decode("utf8") + codecs.BOM_UTF8.decode() + l_encode[0]
    w.writerow(l_encode)
    yield io.getvalue()
    io.seek(0)
    io.truncate(0)
    for c in chat_list:
        try:
            candidate_info = query_candidate_by_id(c[2])
            job_name = job_name
            candidate_id = c[2]
            create_time = c[9].strftime("%Y-%m-%d %H:%M:%S")
            candidate_name = c[3]
            if c[4] == 'user_ask':
                source = '候选人主动'
            elif c[4] == 'search':
                source = '机器人打招呼'
            else:
                source = '未知'
            if c[6] is None:
                wechat = ''
                phone = ''
                resume = ''
            else:
                try:
                    contact = json.loads(c[6])
                    wechat = contact['wechat'] or ''
                    phone = contact['phone'] or ''
                    resume = contact['cv'] or ''
                except Exception as e:
                    wechat = ''
                    phone = ''
                    resume = ''
                    logger.info(
                        f'exception_filter:{candidate_id}, {candidate_name}, {c[6]}, {contact}, {e}, {e.args}, {traceback.format_exc()}')
            try:
                conversation = json.loads(c[7])
                con_str = ''
                for c in conversation:
                    con_str = con_str + c['speaker'] + ':' + c['msg'] + '\n'
            except Exception as e:
                con_str = ''

            if len(candidate_info) == 0:
                logger.info(f"chat_candidate_not_match, {candidate_id}")
                region = ''
                gender = ''
                age = 0
                position = ''
                degree = ''
                major = ''
                large_comps = ''
                school = ''
                edu = ''
                work = ''
                exp_location = ''
                exp_salary = ''
                tag_list = ''
            else:
                c_j = candidate_info[0][7].replace('\n', '.')
                c_j = c_j.replace("\'", '\"')
                candidate_json = json.loads(c_j, strict=False)
                region = candidate_info[0][4]
                if candidate_json.get('gender', 0) == 0:
                    gender = '未知'
                elif candidate_json.get('gender', 0) == 1:
                    gender = '男'
                elif candidate_json.get('gender', 0) == 2:
                    gender = '女'
                else:
                    gender = '未知'
                age = candidate_json.get('age', -1)
                position = candidate_info[0][5]
                if candidate_json.get('degree', -1) == 0:
                    degree = '大专'
                elif candidate_json.get('degree', -1) == 1:
                    degree = '本科'
                elif candidate_json.get('degree', -1) == 2:
                    degree = '硕士'
                elif candidate_json.get('degree', -1) == 3:
                    degree = '博士'
                else:
                    degree = '未知'
                major = candidate_json.get('major', '')
                large_comps = ','.join(candidate_json.get('companies', []))
                school = ''
                if len(candidate_json.get('education', [])) > 0:
                    school = candidate_json['education'][0].get('school', '')
                edu = ''
                for s in candidate_json.get('education', []):
                    edu = edu + s.get('school', '') + ',' + s.get('department', '') + ',' + s.get('sdegree',
                                                                                                  '') + ',' + s.get('v',
                                                                                                                    '') + '\n'
                work = ''
                for e in candidate_json.get('work', []):
                    des = e.get('description', '') or ''
                    work = work + ',' + e.get('company', '') + ',' + e.get('position', '') + ',' + e.get('worktime',
                                                                                                         '') + ',' + e.get(
                        'v', '') + ',' + des + '\n'
                exp_location = candidate_json.get('exp_location', '')
                exp_salary = candidate_json.get('exp_salary', '')
                tag_list = ','.join(candidate_json.get('tag_list', []))

            l = [job_id, candidate_id, create_time, candidate_name, source, wechat, phone, resume, con_str, region,
                 gender, age, position, degree, major, large_comps, school, edu, work, exp_location, exp_salary,
                 tag_list]
            l_encode = [csv_encode(_l) for _l in l]
            w.writerow(l_encode)
            yield io.getvalue()
            io.seek(0)
            io.truncate(0)
        except Exception as e:
            logger.info(f'test_download_candidate_maimai_error4, {c_j}')
            logger.info(
                f'test_download_candidate_maimai_error4,{candidate_id}, {e}, {e.args}, {traceback.format_exc()}')


def pNull(s):
    if s is None or s == None or s == 'null':
        return ''
    return s or ''


def generate_resume_csv_Linkedin(manage_account_id, platform, start_date, end_date, list_name):
    res = get_resume_by_filter(manage_account_id, platform, start_date, end_date, list_name)
    io = StringIO()
    w = csv.writer(io)

    l = ['候选人ID', '创建时间', '候选人姓名', '电话', '邮箱', '地区', '岗位', '最高学历', '专业', '毕业院校', '年龄',
         '教育经历', '工作经历', '语言能力', '工作总结']
    l_encode = [csv_encode(_l) for _l in l]
    l_encode[0] = codecs.BOM_UTF8.decode("utf8") + codecs.BOM_UTF8.decode() + l_encode[0]
    w.writerow(l_encode)
    yield io.getvalue()
    io.seek(0)
    io.truncate(0)
    for r in res:
        try:
            candidate_id = r[1]
            create_time = r[4].strftime("%Y-%m-%d %H:%M:%S")
            profile_json = r[5].replace('\n', ',')
            profile_json = profile_json.replace('/', '_')
            profile_json = profile_json.replace('，', ',')
            profile_json = profile_json.replace('u0000', ',')
            profile_json = profile_json.replace('\\', ',')
            profile = json.loads(profile_json, strict=False).get('profile', {})
            candidate_name = profile.get('name', '')
            phone = ','.join(profile.get('phones', []))
            email = ','.join(profile.get('emails', []))
            region = profile.get('location', '')
            position = profile.get('role', '')
            sdegree = ''
            if len(profile.get('educations', [])) > 0:
                sdegree = profile.get('educations', [])[0].get('degreeInfo', '')
            major = ''
            if len(profile.get('educations', [])) > 0:
                major = profile.get('educations', [])[0].get('majorInfo', '')
            school = ''
            if len(profile.get('educations', [])) > 0:
                school = profile.get('educations', [])[0].get('schoolName', '')
            age = -1
            for i in range(len(profile.get('educations', [])) - 1, -1, -1):
                time_info = profile.get('educations', [])[i].get('timeInfo', '')
                if time_info != '':
                    index = -1
                    for idx, ch in enumerate(time_info):
                        if ch.isdigit():
                            index = idx
                            break
                    if index != -1:
                        age = int(datetime.datetime.today().year) - int(time_info[index: index + 4]) + 18
                        break

            edu = ''
            for e in profile.get('educations', []):
                edu = f"{edu}{pNull(e.get('schoolName', ''))},{pNull(e.get('majorInfo', ''))},{pNull(e.get('degreeInfo', ''))},{pNull(e.get('timeInfo', ''))},{pNull(e.get('summary', ''))}\n"
            work = ''
            for e in profile.get('experiences', []):
                work = f"{work}{pNull(e.get('companyName', ''))},{pNull(e.get('timeInfo', ''))}\n"
                for wo in e.get('works', []):
                    work = f"{work}{pNull(wo.get('workTimeInfo', ''))},{pNull(wo.get('worklocationInfo', ''))},{pNull(wo.get('workPosition', ''))},{pNull(wo.get('workDescription', ''))}\n"

            languages = ''
            for lan in profile.get('languages', []):
                languages = f"{languages}{pNull(lan.get('language', ''))},{pNull(lan.get('des', ''))}\n"
            summary = profile.get('summary', '')

            l = [candidate_id, create_time, candidate_name, phone, email, region, position, sdegree, major, school, age,
                 edu, work, languages, summary]
            l_encode = [csv_encode(_l) for _l in l]
            w.writerow(l_encode)
            yield io.getvalue()
            io.seek(0)
            io.truncate(0)
        except Exception as e:
            logger.info(f'download_resume_error_linkedin,{profile_json}')
            logger.info(f'download_resume_error_linkedin,{candidate_id}, {e}, {e.args}, {traceback.format_exc()}')


def generate_resume_csv_maimai(manage_account_id, platform, start_date, end_date):
    res = get_resume_by_filter(manage_account_id, platform, start_date, end_date)
    io = StringIO()
    w = csv.writer(io)

    l = ['候选人ID', '创建时间', '候选人姓名', '地区', '性别', '年龄', '工作总时长', '在职公司', '岗位', '最高学历',
         '专业', '历史公司', '毕业院校', '教育经历', '工作经历', '预期职位', '预期地点', '预期薪水', '其他倾向',
         '简历标签']
    l_encode = [csv_encode(_l) for _l in l]
    l_encode[0] = codecs.BOM_UTF8.decode("utf8") + codecs.BOM_UTF8.decode() + l_encode[0]
    w.writerow(l_encode)
    yield io.getvalue()
    io.seek(0)
    io.truncate(0)
    for r in res:
        try:
            candidate_id = r[1]
            create_time = r[4].strftime("%Y-%m-%d %H:%M:%S")
            profile_json = r[5].replace('\n', ',')
            profile_json = profile_json.replace('/', '_')
            profile_json = profile_json.replace('，', ',')
            profile_json = profile_json.replace('u0000', ',')
            profile_json = profile_json.replace('\\', ',')
            profile = json.loads(profile_json, strict=False)

            # logger.info(f'download_resume_error_json, {candidate_id}, {e}, {e.args}, {traceback.format_exc()}')
            candidate_name = profile.get('name', '')
            region = profile.get('city', '')
            gender = profile.get('gender_str', '')
            age = profile.get('age', 0)
            work_time = profile.get('work_time', '')
            company = profile.get('company', '')
            position = profile.get('position', '')
            sdegree = profile.get('sdegree', '')
            major = profile.get('major', '')
            large_comps = profile.get('large_comps', '')
            school = profile.get('school', '')
            schools = ''
            for s in profile.get('edu', []):
                schools = schools + s.get('school', '') + ',' + s.get('department', '') + ',' + s.get('sdegree',
                                                                                                      '') + ',' + s.get(
                    'v', '') + '\n'
            if 'current_company' in profile:
                work_detail = profile['current_company'].get('company', '') + ',' + profile['current_company'].get(
                    'position', '') + ',' + profile['current_company'].get('worktime', '') + '\n'
            else:
                work_detail = ''
            for e in profile.get('exp', []):
                des = e.get('description', '') or ''
                work_detail = work_detail + ',' + e.get('company', '') + ',' + e.get('position', '') + ',' + e.get(
                    'worktime', '') + ',' + e.get('v', '') + ',' + des + '\n'

            exp_positon = ','.join(profile['job_preferences'].get('positons', []))
            exp_location = ','.join(profile['job_preferences'].get('province_cities', []))
            exp_salary = profile['job_preferences'].get('salary', '')
            exp_prefer = ','.join(profile['job_preferences'].get('prefessions', []))
            tags = ','.join(profile.get('tag_list', []))
            l = [candidate_id, create_time, candidate_name, region, gender, age, work_time, company, position, sdegree,
                 major, large_comps, school, schools, work_detail, exp_positon, exp_location, exp_salary, exp_prefer,
                 tags]
            l_encode = [csv_encode(_l) for _l in l]
            w.writerow(l_encode)
            yield io.getvalue()
            io.seek(0)
            io.truncate(0)
        except Exception as e:
            logger.info(f'download_resume_error4,{profile_json}')
            logger.info(f'download_resume_error4,{candidate_id}, {e}, {e.args}, {traceback.format_exc()}')


def format_json_str(body):
    r = ""
    if type(body) == list:
        for sub in body:
            r += format_json_str(sub)
            r += "\n"
    elif type(body) == dict:
        for k in body:
            r += k
            r += ":"
            r += format_json_str(body[k])
            r += '\t'
    elif type(body) == str:
        r = body[1:-1] if len(body) > 0 and body[0] == "\"" and body[-1] == "\"" else body
    return r


def fetch_json_str(body, key, default_v=''):
    return format_json_str(body[key]) if key in body else default_v


def generate_csv(res):
    s = res[0][6].replace('\n', '\\n')
    res_l = json.loads(s)
    format_resume_json = json.loads(res[0][8].replace('\n', '\\n'))
    io = StringIO()
    w = csv.writer(io)
    w.writerow(['简历', '结果', '匹配结果', '姓名', '性别', '年龄/出生', '期望职位', '期望薪资', '最高学历', '专业',
                '工作经历', '教育经历', '工作城市', '电话', '邮箱', '技能', '项目经历'])
    for idx, r in enumerate(res_l):
        file_name = os.path.basename(r['f_path'])
        l = [file_name, r['res'], r['remark'],
             fetch_json_str(format_resume_json[idx], '姓名'), fetch_json_str(format_resume_json[idx], '性别'),
             fetch_json_str(format_resume_json[idx], '年龄/出生'), fetch_json_str(format_resume_json[idx], '期望职位'),
             fetch_json_str(format_resume_json[idx], '期望薪资'), fetch_json_str(format_resume_json[idx], '最高学历'),
             fetch_json_str(format_resume_json[idx], '专业'), fetch_json_str(format_resume_json[idx], '工作经历'),
             fetch_json_str(format_resume_json[idx], '教育经历'), fetch_json_str(format_resume_json[idx], '工作城市'),
             fetch_json_str(format_resume_json[idx], '电话'), fetch_json_str(format_resume_json[idx], '邮箱'),
             fetch_json_str(format_resume_json[idx], '技能'), fetch_json_str(format_resume_json[idx], '项目经历')]
        w.writerow(l)
        yield io.getvalue()
        io.seek(0)
        io.truncate(0)


def imglist_to_text(img_url_list):
    res = ''
    for img in img_url_list:
        res = res + img_to_text(img)
    return res


def img_to_text(f_path):
    res = reader.readtext(f_path)
    s = ''
    for r in res:
        s = s + r[1] + '\n'
    return s


def download_file(url):
    file_name = os.path.basename(url)
    responsepdf = requests.get(url)
    if responsepdf.status_code == 200:
        dst_path = os.path.join(file_path_prefix, file_name)
        with open(dst_path, "wb") as f:
            f.write(responsepdf.content)
        return dst_path
    else:
        return None


def content_transfer(f_path):
    file_name = os.path.basename(f_path)
    if f_path.endswith('jpg') or f_path.endswith('jpeg') or f_path.endswith('png'):
        return True, img_to_text(f_path)
    elif f_path.endswith('pdf'):
        final_url_list = []
        start_time = datetime.datetime.now()
        pdf = fitz.open(f_path)
        for page in pdf:
            mat = fitz.Matrix(10, 10)
            pix = page.get_pixmap(matrix=mat)
            png_file_name = f"{file_name.split('.')[0]}-page-{page.number}.png"
            pix.save(png_file_name)
            final_url_list.append(png_file_name)
        pdf.close()
        end_pdf_time = datetime.datetime.now()
        res = imglist_to_text(final_url_list)
        end_txt_time = datetime.datetime.now()
        logger.info(
            f'[tool_service] content transfer pdf -> jpg time consumption = {(end_pdf_time - start_time).total_seconds()}, jpg -> text time consumption = {(end_txt_time - end_pdf_time).total_seconds()}')
        for u in final_url_list:
            os.remove(u)
        return True, res
    elif f_path.endswith('docx'):
        f = docx.Document(f_path)
        result_str = ''
        for p in f.paragraphs:
            if p.text.strip() != '':
                result_str += ';' + p.text
        tables = f.tables
        last_cell = ''
        for t in tables:
            for r in t.rows:
                for c in r.cells:
                    s = c.text.strip()
                    if s != '' and s != last_cell:
                        result_str += ';' + s
                        last_cell = s
        return True, result_str
    elif f_path.endswith('doc'):
        return False, "EXT_NOT_SUPPORT"
    else:
        return False, "EXT_NOT_SUPPORT"


def content_extract_and_filter(file_raw_data, jd):
    # chatgpt = ChatGPT()
    # 先做个人为截断，如果有问题再说
    start_time = datetime.datetime.now()
    ext_prompt_msg = '以下是候选人信息，请提取关键信息并结构化输出\n$$$\n' + file_raw_data[0:3500] + "\n$$$"
    ext_prompt = Prompt()
    ext_prompt.add_user_message(ext_prompt_msg)
    file_key_data = gpt_manager.chat_task(ext_prompt)
    end_format_time = datetime.datetime.now()
    # file_key_data = chatgpt.chat(ext_prompt)
    extract_prompt_msg = '你是一个猎头, 请从候选人信息提取: 姓名, 性别, 年龄/出生, 期望职位, 期望薪资, 最高学历, 专业, 教育经历, 工作经历, 工作城市, 电话, 邮箱, 技能, 项目经历.\n如果候选人信息不包含该字段, 标记为空, 请用中文回答, 以json格式输出'
    extract_prompt_msg += f'$$$\n候选人个人信息如下：{file_key_data}\n$$$\n'
    extract_prompt = Prompt()
    extract_prompt.add_user_message(extract_prompt_msg)
    extract_info = gpt_manager.chat_task(extract_prompt)
    format_info = {}
    try:
        if '```json' in extract_info:
            extract_info = extract_info.replace('```json', '')
        if '```' in extract_info:
            extract_info = extract_info.replace('```', '')
        format_info = json5.loads(extract_info)
    except BaseException:
        pass
    end_parse_time = datetime.datetime.now()
    logger.info(f"filter_task_content_extract_and_filter_file_key_data: {file_key_data}")
    prefix = '你是一个猎头，请判断候选人是否符合招聘要求\n给出具体原因和推理过程\n答案必须在最后一行，并且单独一行 A.合适，B.不合适'
    candidate_msg = f'$$$\n候选人个人信息如下：{file_key_data}\n$$$\n'
    filter_prompt_msg = prefix + candidate_msg + '\n招聘要求:\n' + jd + '\n'
    filter_prompt = Prompt()
    filter_prompt.add_user_message(filter_prompt_msg)
    res = gpt_manager.chat_task(filter_prompt)
    end_judge_time = datetime.datetime.now()

    logger.info(
        f'[content_extract_and_filter] format resume time = {(end_format_time - start_time).total_seconds()}, parse result time = {(end_parse_time - end_format_time).total_seconds()}, judge time = {(end_judge_time - end_parse_time).total_seconds()}')
    # res = chatgpt.chat(filter_prompt)
    return res, format_info


def exec_filter_task(manage_account_id, file_list, jd):
    filter_result = []
    format_resume_infos = []
    for f_path in file_list:
        flag, file_raw_data = content_transfer(f_path)
        logger.info(f"filter_task_content_transfer:{f_path}, {flag}, {len(file_raw_data)}, {file_raw_data[0:3500]}")
        file_name = basename(f_path)
        if not flag:
            logger.info(f'file_ext_not_support, {manage_account_id}, {f_path}')
            filter_result.append({
                "f_path": file_name,
                "res": file_raw_data,
                "remark": ""
            })
            continue

        single_filter_result, format_resume_info = content_extract_and_filter(file_raw_data, jd)
        logger.info(f"filter_task_content_extract_and_filter:{f_path}, {single_filter_result}, {format_resume_info}")
        res = 'QUALIFIED' if 'A.合适' in single_filter_result else 'UNQUALIFIED'
        filter_result.append({
            "f_path": file_name,
            "res": res,
            "remark": single_filter_result
        })
        format_resume_infos.append(format_resume_info)
    return filter_result, format_resume_infos


def customized_user_scenario(manage_account_id, context, platform, scenario_info, extra_info=''):
    create_customized_scenario_setting(manage_account_id, platform, context, scenario_info, extra_info='')


def update_user_scenario(rid, scenario_info, extra_info=''):
    update_customized_scenario_setting(scenario_info, extra_info, rid)


def create_customized_greeting_service(manage_account_id, platform, scenario_info):
    create_customized_greeting(manage_account_id, platform, scenario_info)


def update_customized_greeting_service(scenario_info, rid):
    update_customized_greeting(scenario_info, rid)


def delete_customized_greeting_service(rid):
    delete_customized_greeting(rid)


def query_customized_greeting_service(manage_account_id, platform):
    query_customized_greeting(manage_account_id, platform)


def get_email_template(manage_account_id, platform):
    scenario_info = query_customized_scenario_setting(manage_account_id, platform, SCENARIO_EMAIL)
    if len(scenario_info) == 0 or len(scenario_info[0]) == 0:
        return {}
    scenario_info = scenario_info[0][0]
    if scenario_info == None or len(scenario_info) == 0:
        return {}
    else:
        return json.loads(scenario_info, strict=False)


def get_inmail_template(manage_account_id, platform):
    scenario_info = query_customized_scenario_setting(manage_account_id, platform, SCENARIO_INMAIL)
    if len(scenario_info) == 0 or len(scenario_info[0]) == 0:
        return {}
    scenario_info = scenario_info[0][0]
    if scenario_info == None or len(scenario_info) == 0:
        return {}
    else:
        return json.loads(scenario_info, strict=False)


def get_default_email_template(idx, platform):
    total = get_default_email_template_count(platform)[0][0]
    template = get_default_email_template_by_idx(platform, idx)
    if len(template) == 0 or len(template[0]) == 0:
        return None, f'total template is {total}, {idx} is exceeded'
    template = template[0][0]
    return {'total': total, 'idx': idx, 'template': template}, None


def get_default_inmail_template(idx, platform):
    total = get_default_inmail_template_count(platform)[0][0]
    template = get_default_inmail_template_by_idx(platform, idx)
    if len(template) == 0 or len(template[0]) == 0:
        return None, f'total template is {total}, {idx} is exceeded'
    template = template[0][0]
    return {'total': total, 'idx': idx, 'template': template}, None


def get_default_greeting_template_v2(platform):
    templates = get_all_default_greeting_template(platform)
    if len(templates) == 0 or len(templates[0]) == 0:
        return None, f'no greeting template'
    res = []
    for template in templates:
        res.append(template[0])
    return res, None


def get_default_greeting_template(idx, platform):
    total = get_default_greeting_template_count(platform)[0][0]
    template = get_default_greeting_template_by_idx(platform, idx)
    if len(template) == 0 or len(template[0]) == 0:
        return None, f'total template is {total}, {idx} is exceeded'
    template = template[0][0]
    return {'total': total, 'idx': idx, 'template': template}, None


def flush_email_credentials(manage_account_id, email, pwd, platform):
    logger.info(f'[flush_email_credentials] {manage_account_id} {email} {pwd} {platform}')
    flush_email_credentials_db(manage_account_id, email, pwd, platform)


def get_email_credentials(manage_account_id, platform):
    email_credential = get_email_credentials_db(manage_account_id, platform)
    logger.info(f'[get_email_credentials] {manage_account_id} {platform} -> {email_credential}')
    if len(email_credential) == 0 or len(email_credential[0]) == 0:
        return None, f'{manage_account_id}没有注册邮箱配置'
    email_credential = email_credential[0]
    return {'manage_account_id': email_credential[0], 'email': email_credential[1], 'pwd': email_credential[2],
            'platform': email_credential[3]}, None


def get_default_greeting_scenario():
    msg = 'Hi \n'
    msg += 'we are looking for an candidate base in Irvine/Seattle for FFALCON who is expanding streaming business, it\'s the leading smart TVs & AIoT company in China\n'
    msg += 'your Exp. seems a good match\n'
    msg += 'would you like to explore this opportunity? Thanks!'
    return {'默认': msg}


def get_default_greeting_scenario_v2():
    msg = 'Hi \n'
    msg += 'we are looking for an candidate base in Irvine/Seattle for FFALCON who is expanding streaming business, it\'s the leading smart TVs & AIoT company in China\n'
    msg += 'your Exp. seems a good match\n'
    msg += 'would you like to explore this opportunity? Thanks!'
    return [{'id': 0, 'msg': msg}]


def get_default_chat_scenario():
    scenario_options = ['要简历', '约电话', '转介绍', '召回']
    r = {}
    for scenario in scenario_options:
        msg = 'Hi, \n'
        if scenario == '要简历':
            msg += 'Thanks for getting back to me! Could you please send me your updated resume so we can move forward to the next step? Also, can I have your availability for a short call to discuss this opportunity further.\nLooking forward to hearing from you soon. Thanks!'
        if scenario == '约电话':
            msg += 'Thanks for the updated resume and I\'ve well received! Meanwhile, can I have your availability for a short call to discuss this opportunity further.\nLooking forward to hearing from you soon. Thanks!'
        if scenario == '转介绍':
            msg += 'Thanks for the reply, I understand you\'re not actively looking for a new job right now. However, I\'d greatly appreciate your insights. Do you happen to know someone in your network who might be interested in this role? Feel free to pass along the details, and if they have questions, they can reach out directly.\nThanks again for any help you can provide!'
        if scenario == '召回':
            msg += 'Thanks for connecting! I trust this message finds you in good spirits. I noticed your noteworthy background on LinkedIn!\nCurrently, FFALCON is on a global expansion drive. It\'s the sub-brand of TCL Electronics, an established global TV manufacturing brand. To strengthen TCL Corporation’s globalization strategy plans within the smart home sector, FFalcon has been developed into a leading brand with a business value of over 650 million USD.\nI believe your insights could significantly contribute to our strategy, and I would like to discuss more details about this opportunity with you! Would you mind sharing a concise update on your CV? Your expertise aligns with our objectives, and your input would be immensely valuable.'
        r[scenario] = msg
    return r


def get_leave_msg(manage_account_id, platform):
    scenario_info = query_customized_scenario_setting(manage_account_id, platform, SCENARIO_GREETING)
    if len(scenario_info) == 0 or len(scenario_info[0]) == 0:
        return get_default_greeting_scenario()
    scenario_info = scenario_info[0][0]
    if scenario_info == None or len(scenario_info) == 0:
        return get_default_greeting_scenario()
    else:
        return json.loads(scenario_info, strict=False)


def get_leave_msg_v2(manage_account_id, platform):
    scenario_info = query_customized_greeting(manage_account_id, platform)
    if len(scenario_info) == 0 or len(scenario_info[0]) == 0:
        return get_default_greeting_scenario_v2()
    else:
        ret = []
        logger.info("get_leave_msg_v2 = {}".format(scenario_info))
        for rid, s_info in scenario_info:
            ret.append({'id': rid, 'msg': s_info})
        return ret


def get_chat_scenario(manage_account_id, platform):
    scenario_info = query_customized_scenario_setting(manage_account_id, platform, SCENARIO_CHAT)
    if len(scenario_info) == 0 or len(scenario_info[0]) == 0:
        return get_default_greeting_scenario()
    scenario_info = scenario_info[0][0]
    if scenario_info == None or len(scenario_info) == 0:
        return get_default_greeting_scenario()
    else:
        return json.loads(scenario_info, strict=False)


_tag_id_cache = {}
_tag_name_id_cache = {}


def ensure_cache(manage_account_id, platform):
    if manage_account_id not in _tag_id_cache:
        _tag_id_cache[manage_account_id] = {}
        _tag_name_id_cache[manage_account_id] = {}
        id_tags = query_profile_id_tag(manage_account_id, platform)
        for id_tag in id_tags:
            _tag_id_cache[manage_account_id][id_tag[0]] = id_tag[1]
            _tag_name_id_cache[manage_account_id][id_tag[1]] = id_tag[0]
    return _tag_id_cache[manage_account_id], _tag_name_id_cache[manage_account_id]


def get_check_tag_ids(manage_account_id, tags, platform):
    user_id_tag_cache, user_tag_id_cache = ensure_cache(manage_account_id, platform)
    tag_ids = []
    for tag in tags:
        if tag not in user_tag_id_cache:
            return None
        else:
            tag_ids.append(user_tag_id_cache[tag])
    return tag_ids


def create_profile_tag(manage_account_id, platform, tag):
    user_id_tag_cache, user_tag_id_cache = ensure_cache(manage_account_id, platform)
    if tag in user_tag_id_cache:
        logger.info("[tool_service] tag = {} already exist, ignore create")
        return {'tag_id': user_tag_id_cache[tag], 'tag': tag, 'manage_account_id': manage_account_id,
                'platform': platform, 'tag': tag}, None
    try:
        tag_id = create_profile_tag_db(manage_account_id, platform, tag)
    except BaseException as e:
        logger.error("[tool_service] create profile tag failed for {}, {}, {}".format(manage_account_id, platform, tag))
        return None, "创建profile tag失败"
    user_id_tag_cache[tag_id] = tag
    user_tag_id_cache[tag] = tag_id

    return {'tag_id': tag_id, 'tag': tag, 'manage_account_id': manage_account_id, 'platform': platform,
            'tag': tag}, None


def query_profile_tag_by_user(manage_account_id, platform):
    user_id_tag_cache, user_tag_id_cache = ensure_cache(manage_account_id, platform)
    tags = list(user_tag_id_cache.keys())
    return tags, None


def query_profile_tag_relation_by_user_and_candidate(manage_account_id, candidate_id, platform):
    id_tags = query_profile_tag_relation_by_user_and_candidate_db(manage_account_id, candidate_id, platform)
    tags = []
    for id_tag in id_tags:
        tags.append(id_tag[1])

    return tags, None


def associate_profile_tags(manage_account_id, candidate_id, platform, tags):
    tag_ids = get_check_tag_ids(manage_account_id, tags, platform)
    if not tag_ids:
        return None, "tags中存在无效tag"
    logger.info("[tool_service] associate_profile_tags tag_ids = {}, tags = {}", tag_ids, tags)
    for idx, tag_id in enumerate(tag_ids):
        relations = query_id_by_profile_tag_relation(manage_account_id, candidate_id, platform, [tags[idx]])
        if relations and len(relations) > 0:
            logger.info(
                "[tool_service] associate_profile_tag existing relation manage_account_id = {}, candidate_id = {}, platform = {}, tag_id = {}, tag = {}",
                manage_account_id, candidate_id, platform, tag_id, tags[idx])
            continue
        associate_profile_tag(manage_account_id, candidate_id, platform, tag_id, tags[idx])
        logger.info(
            "[tool_service] associate_profile_tag manage_account_id = {}, candidate_id = {}, platform = {}, tag_id = {}, tag = {}",
            manage_account_id, candidate_id, platform, tag_id, tags[idx])
    return tags, None


def upload_profile_status(manage_account_id, candidate_id, platform, profile):
    if "status" in profile:
        status = profile["status"]
    else:
        status = None

    if status is None:
        return

    upload_profile_status_dao(manage_account_id, candidate_id, platform, status)


def deassociate_profile_tags(manage_account_id, candidate_id, platform, tags):
    tag_ids = get_check_tag_ids(manage_account_id, tags, platform)
    if not tag_ids:
        return None, "tags中存在无效tag"
    delete_profile_tag_relation(manage_account_id, candidate_id, platform, tags)
    logger.info(
        "[tool_service] delete_profile_tag_relation manage_account_id = {}, candidate_id = {}, platform = {}, tag_ids = {}, tags = {}",
        manage_account_id, candidate_id, platform, tag_ids, tags)
    return tags, None


def delete_profile_tags(manage_account_id, candidate_id, platform, tags):
    user_id_tag_cache, user_tag_id_cache = ensure_cache(manage_account_id, platform)
    tag_ids = get_check_tag_ids(manage_account_id, tags, platform)
    if not tag_ids:
        return None, "tags中存在无效tag"
    relations = query_id_by_profile_tag_relation(manage_account_id, candidate_id, platform, tags)
    if relations and len(relations) != 0:
        return None, "存在标记的tag关系, 请先清除profile和tag关系"
    delete_profile_tags_db(tag_ids)
    for idx, tag in enumerate(tags):
        user_id_tag_cache.pop(tag_ids[idx])
        user_tag_id_cache.pop(tag)
    logger.info("[tool_service] delete manage_account_id = {}, candidate_id = {}, platform = {}, tags = {}",
                manage_account_id, candidate_id, platform, tags)
    return None, None


def transfer_profile():
    # to do refactor for download excel
    return None


def get_max_time_info(time_info_str, default_time):
    if not time_info_str:
        return default_time
    max_start_year = default_time
    times = re.findall(r'\d\d\d\d', time_info_str)
    for t in times:
        max_start_year = max(max_start_year, int(t))
    return max_start_year


def get_min_time_info(time_info_str, default_time):
    if not time_info_str:
        return default_time
    min_start_year = default_time
    times = re.findall(r'\d\d\d\d', time_info_str)
    for t in times:
        min_start_year = min(min_start_year, int(t))
    # logger.info(f"get_min_time_info time_info_str: {time_info_str} times: {times} min_start_year: {min_start_year}")
    return min_start_year


def search_profile_by_tag(manage_account_id, platform, tags, page, limit, contact2str):
    tag_ids = get_check_tag_ids(manage_account_id, tags, platform)
    if not tag_ids:
        return None, "tags中存在无效tag"
    candidate_ids = query_candidate_id_by_tag_relation(manage_account_id, platform, tags)
    total_count = get_resume_total_count_by_candidate_ids_and_platform(manage_account_id, platform, candidate_ids)
    start = (page - 1) * limit
    rows = get_resume_by_candidate_ids_and_platform(manage_account_id, platform, candidate_ids, start, limit)
    details = []
    data = {'page': page, 'limit': limit, 'total': total_count, 'details': details}

    for row in rows:
        profile = parse_profile(row[1], 'need_deserialize', contact2str)
        if profile is None:
            continue
        profile['candidateId'] = row[0]
        profile['cvUrl'] = row[2]
        details.append(profile)
    return data, None


def fetch_contact_infos(manage_account_id, candidate_ids):
    ret_dict = {}
    if not candidate_ids:
        return ret_dict
    b = time.time()
    user_links = query_user_link_by_id_set(manage_account_id, candidate_ids)
    contacts = query_contact_by_id_set(linkedin_id_set=candidate_ids)
    contact_dict = {}
    for linkedin_id, personal_email, work_email, phone in contacts:
        contact_dict[linkedin_id] = {
            'Email': json.loads(personal_email) + json.loads(work_email),
            'Phone': json.loads(phone)
        }
    for link_linkedin_id, link_contact_type in user_links:
        if link_linkedin_id not in ret_dict:
            ret_dict[link_linkedin_id] = {}
        contact_key = 'Email' if 'email' in link_contact_type else 'Phone'
        if contact_key not in ret_dict[link_linkedin_id]:
            ret_dict[link_linkedin_id][contact_key] = []
        contact_content = contact_dict.get(link_linkedin_id, {}).get(contact_key, [])
        ret_dict[link_linkedin_id][contact_key] += contact_content
    e = time.time()
    logger.info(f'fetch_contact_infos, from |{contacts}| |{user_links}| to |{ret_dict}|, cost: {e - b}s')
    return ret_dict


def search_profile_by_tag_v2(manage_account_id, platform, tag, company, candidate_name, status, stage, page, limit,
                             contact2str):
    total_count = query_tag_filter_num_new(manage_account_id, platform, tag, company, candidate_name, stage, status)
    start = (page - 1) * limit

    details = []
    data = {'page': page, 'limit': limit, 'total': total_count, 'details': details}
    rows = query_tag_filter_profiles_new(manage_account_id, platform, tag, company, candidate_name, stage, status,
                                         start, limit)
    candidate_ids = [row[0] for row in rows]
    candidate_contact_infos = fetch_contact_infos(manage_account_id, candidate_ids)
    for row in rows:
        profile = parse_profile(row[1], 'need_deserialize', contact2str)
        candidate_id = row[0]
        if candidate_id in candidate_contact_infos:
            if not profile['contactInfo'].get('Phone', ''):
                bank_phone = candidate_contact_infos[candidate_id].get('Phone', [])
                if len(bank_phone):
                    logger.info(
                        f'manage_account_id {manage_account_id}, linkedin {candidate_id}, will update phone to: ({len(bank_phone)}) {bank_phone[0]}')
                    profile['contactInfo']['Phone'] = bank_phone[0]
            if not profile['contactInfo'].get('Email', ''):
                bank_email = candidate_contact_infos[candidate_id].get('Email', [])
                if len(bank_email):
                    logger.info(
                        f'manage_account_id {manage_account_id}, linkedin {candidate_id}, will update email to: ({len(bank_email)}) {bank_email[0]}')
                    profile['contactInfo']['Email'] = bank_email[0]
        if profile is None:
            continue
        profile['candidateId'] = row[0]
        profile['cvUrl'] = row[2]
        profile['status'] = row[3]
        profile['abstract'] = fetch_abstract(profile)
        # stage = query_stage_by_id(manage_account_id, platform, tag, profile['candidateId'])
        profile["stage"] = row[4]
        profile['experiences'] = None
        details.append(profile)
    return data, None


def fetch_abstract(profile):
    if "experiences" not in profile:
        return ""

    abstract = ""
    for experience in profile["experiences"][:3]:
        company = experience["companyName"] if "companyName" in experience else ""
        title = experience["works"][0]["workPosition"] if len(experience["works"]) > 0 and "workPosition" in \
                                                          experience["works"][0] else ""
        time_info = experience["works"][0]["workTimeInfo"] if len(experience["works"]) > 0 and "workTimeInfo" in \
                                                              experience["works"][0] else ""
        abstract += f"{company} {title} {time_info}\n"

    return abstract


def generate_email_content(manage_account_id, platform, candidate_id, template):
    email_template = get_email_template(manage_account_id, platform)
    if email_template == {}:
        return None, f'{candidate_id}未设置邮件模板'
    if template not in email_template:
        return None, f'{template}不存在, {list(email_template.keys())}'
    template_val = email_template[template]
    rows = get_resume_by_candidate_ids_and_platform(manage_account_id, platform, [candidate_id], 0, 10)
    if len(rows) == 0 or len(rows[0]) == 0:
        return template_val, None
    profile = parse_profile(rows[0][1], 'need_deserialize', True)
    for key in profile:
        template_key = '#{' + key + '}'
        if profile[key] and template_key in template_val:
            template_val = template_val.replace(template_key, profile[key])
    return template_val, None


def send_email_163(email_from, email_to, pwd, subject, body):
    email_user = email_from
    email_password = pwd
    to_emails = [email_to]

    subject = subject
    body = body

    msg = MIMEMultipart()
    msg['From'] = email_user
    msg['To'] = ', '.join(to_emails)
    msg['Subject'] = subject

    msg.attach(MIMEText(body, 'plain'))

    all_recipients = to_emails
    ret = False
    try:
        server = smtplib.SMTP_SSL('smtp.163.com', 465)
        server.login(email_user, email_password)
        text = msg.as_string()
        server.sendmail(email_user, all_recipients, text)
        logger.info(f"[send_email_163] {email_from} -> {email_to}, {subject} send successfully")
        ret = True
    except Exception as e:
        logger.error(f"[send_email_163] {email_from} -> {email_to}, {subject} send failed")
        logger.error(traceback.format_exc())
        ret = False
    finally:
        server.quit()
    return ret


def send_email_gmail(email_from, email_to, pwd, subject, body):
    email_user = email_from
    email_password = pwd
    # List of recipients
    to_emails = [email_to]

    # Email content
    subject = subject
    body = body

    # Set up the MIME
    msg = MIMEMultipart()
    msg['From'] = email_user
    msg['To'] = ', '.join(to_emails)
    msg['Subject'] = subject
    msg.attach(MIMEText(body, 'plain'))
    # Connect to the Gmail SMTP server using SMTP_SSL
    server = smtplib.SMTP_SSL('smtp.gmail.com', 465)
    try:
        # Log in to the server
        server.login(email_user, email_password)
        # Send the email
        text = msg.as_string()
        server.sendmail(email_user, to_emails, text)
        logger.info(f"[send_email_gmail] {email_from} -> {email_to}, {subject} send successfully")
        ret = True
    except Exception as e:
        logger.error(f"[send_email_gmail] {email_from} -> {email_to}, {subject} send failed")
        logger.error(traceback.format_exc())
        ret = False
    finally:
        # Quit the server
        server.quit()
    return ret


def send_email_content(manage_account_id, platform, candidate_id, title, content, email_to):
    email_info = query_email_info(manage_account_id)
    if len(email_info) == 0 or len(email_info[0]) == 0:
        return None, f'{manage_account_id}未设置邮箱'
    email_info = email_info[0]
    email_from = email_info[0]
    email_pwd = email_info[1]
    if email_to is None:
        rows = get_resume_by_candidate_ids_and_platform(manage_account_id, platform, [candidate_id], 0, 10)
        if len(rows) == 0:
            return None, f'{candidate_id} 无对应记录'
        profile = deserialize_raw_profile(rows[0][1])
        if profile and 'profile' in profile:
            profile = profile['profile']
        logger.info('[send_email] profile = {} , candidate_id = {}'.format(profile, candidate_id))
        if not profile or 'contactInfo' not in profile or 'Email' not in profile['contactInfo'] or len(
                profile['contactInfo']['Email']) == 0:
            return None, f'{candidate_id} 无email联系方式'
        email_to = profile['contactInfo']['Email']
    # email_to = 'db24@outlook.com'
    logger.info('[send_email] {} {}'.format(email_from, email_to))
    send_ret = False
    if '@163.com' in email_from:
        send_ret = send_email_163(email_from, email_to, email_pwd, title, content)
    elif '@gmail' in email_from:
        send_ret = send_email_gmail(email_from, email_to, email_pwd, title, content)
    else:
        return None, f'不支持邮箱{email_from}'
    if send_ret:
        return None, None
    else:
        return None, f'{email_from}发送失败, 请确认密码正确, enable smtp&pop3设置'


def search_tag_flow_infos(manage_account_id, platform, tag):
    data = query_tag_flow_status(manage_account_id, platform, tag)
    flow_infos = {}
    for d in data:
        flow_status = d[1]
        if flow_status not in flow_infos.keys():
            flow_infos[flow_status] = 0
        flow_infos[flow_status] += 1

    status_infos = {}
    status_infos["connected"] = 0
    status_infos["wait connect"] = 0
    status_infos["pending"] = 0
    companys_dict = {}
    data = query_tag_resume_infos(manage_account_id, platform, tag)
    for d in data:
        company = d[0]
        companys_dict[company] = 1
        status = d[1]
        status_infos[status] += 1

    infos = {
        "flow_infos": flow_infos,
        "status_infos": status_infos,
        "companys": list(companys_dict)
    }

    return infos


def change_flow_status_service(manage_account_id, platform, tag, candidate_id, flow_status):
    update_flow_status(manage_account_id, platform, tag, candidate_id, flow_status)


def get_log(manage_account_id, platform, tag, candidate_id):
    log = fetch_tag_log(manage_account_id, platform, tag, candidate_id)
    return log


def add_tag_log(manage_account_id, platform, tag, candidate_id, flow_status, new_log):
    current_time = datetime.datetime.now()
    now_time = current_time.strftime("%Y-%m-%d %H:%M:%S")
    logs = fetch_tag_log(manage_account_id, platform, tag, candidate_id)
    logs.append(
        {
            "time": now_time,
            "flow_status": flow_status,
            "msg": new_log
        }
    )
    logger.info(
        f"add_tag_log manage_account_id: {manage_account_id} platform: {platform} tag: {tag} candidate_id: {candidate_id} flow_status: {flow_status} logs: {logs}")
    update_tag_log(manage_account_id, platform, tag, candidate_id, json.dumps(logs, ensure_ascii=False))


def parse_profile_gpt(profile):
    prompt_msg = f"{json.dumps(profile)[0:3500]} \nThis is a LinkedIn resume of a candidate. As a headhunter, you need to analyze the resume and summarize two aspects:\n1 => Industry Experience and Expertise\n2 => Career Highlights\nThe result should be represented in JSON format with 'industry_experience' for industry experience and expertise, and 'career_highlights' for career highlights. \nSummarize the content in no more than 50 words, and ensure it is within 50 words.\nThe result content is returned in Chinese."
    prompt = Prompt()
    prompt.add_user_message(prompt_msg)
    output = gpt_manager.chat_task(prompt)
    logger.info(f"output: {output}")
    try:
        details = json.loads(output, strict=False)
    except BaseException as e:
        logger.error(f"parse_profile_gpt json格式异常: {output}")
        return {}

    return details


def parse_profile_by_ai_service(manage_account_id, platform, candidate_id, use_ai):
    rows = get_resume_by_candidate_ids_and_platform(manage_account_id, platform, [candidate_id], 0, 10)
    if len(rows) == 0:
        return {}

    raw_profile = deserialize_raw_profile(rows[0][1])
    profile = parse_profile(raw_profile, 'no', True)
    if 'cv' in profile:
        del profile['cv']

    if use_ai:
        apt_profile_info = parse_profile_gpt(profile)
        if "industry_experience" in apt_profile_info:
            profile["industry_experience"] = apt_profile_info["industry_experience"]
        if "career_highlights" in apt_profile_info:
            profile["career_highlights"] = apt_profile_info["career_highlights"]

    if "workTime" in profile and profile["workTime"] is not None:
        profile["workTimeStr"] = f"{profile['workTime']}年至今"

    if "age" in profile and profile["age"] is not None:
        if profile["age"] <= 0 or profile["age"] > 100:
            profile["age"] = None

    if "last5Jump" in profile and profile["last5Jump"] is not None:
        profile["last5JumpStr"] = f"5年{profile['last5Jump']}跳"

    if "experiences" in profile:
        del profile["experiences"]

    return profile
